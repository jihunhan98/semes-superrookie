"""docx 업로드 확인 데모 — 백엔드 (FastAPI).

목적: 보안 검토 전에 "docx 파일이 웹으로 잘 업로드되고, 내용이 제대로 파싱되는가"만
확인하는 최소 도구. 업로드된 .docx를 python-docx로 파싱해 문단/표를 그대로 웹에 뿌린다.
(모호성 판정 등은 하지 않음 — 업로드·파싱 확인 전용)
"""
import io
import os

from docx import Document
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import FileResponse

HERE = os.path.dirname(os.path.abspath(__file__))
app = FastAPI(title="docx 업로드 확인")


@app.get("/")
def index():
    return FileResponse(os.path.join(HERE, "static", "index.html"))


@app.post("/api/upload")
async def upload(file: UploadFile = File(...)):
    name = file.filename or ""
    data = await file.read()
    if not data:
        return {"ok": False, "error": "빈 파일입니다.", "filename": name}

    # 파일 형식 감지 (.docx 는 실제로 zip = 'PK'로 시작)
    head = data[:8]
    head_hex = " ".join(f"{b:02X}" for b in head)
    if head[:2] != b"PK":
        if head[:4] == b"\xd0\xcf\x11\xe0":  # OLE 복합문서
            hint = ("구버전 Word(.doc) / 한글(.hwp) / 암호·DRM(문서보안) 걸린 docx 로 보입니다. "
                    "Word에선 열려도 실제 파일은 암호화돼 zip이 아닙니다. "
                    "DRM이면 보안 해제된 사본으로, .doc/.hwp면 .docx로 변환해 올리세요.")
        else:
            hint = "확장자만 .docx 이고 실제 내용은 다른 형식일 수 있습니다. (원본 형식/보안 상태 확인)"
        return {"ok": False,
                "error": f"'{name}' 은(는) docx(zip) 형식이 아닙니다. {hint} [파일 시그니처: {head_hex}]",
                "filename": name, "head_hex": head_hex}

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        return {"ok": False,
                "error": f"zip이지만 docx로 열리지 않습니다: {e} (한글 hwpx 등 다른 zip 형식일 수 있음)",
                "filename": name}

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        tables.append([[c.text for c in row.cells] for row in t.rows])

    return {
        "ok": True,
        "filename": name,
        "size_kb": round(len(data) / 1024, 1),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }
